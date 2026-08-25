from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\PI\Projects\PROP INTELLIGENCE\prop_intelligence")
OUT = ROOT / "operations" / "PI_Prop_Intelligence_Operating_Costs_and_Renewal_Calendar_2026.docx"

NAVY = "071521"
NAVY_2 = "0C1D2B"
GOLD = "B59D52"
PALE_GOLD = "F2E8C9"
TEAL = "45D6AC"
WHITE = "FFFFFF"
LIGHT = "E8EDF2"
MID = "AAB4BE"
RED = "C95757"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def margins(cell, top=70, start=85, bottom=70, end=85):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for edge, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, color=NAVY, bold=False, size=8.2, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths, font_size=7.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(width)
        shade(cell, NAVY)
        set_cell_text(cell, header, WHITE, True, 7.4)
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, (value, width) in enumerate(zip(row, widths)):
            cells[idx].width = Inches(width)
            shade(cells[idx], WHITE if row_idx % 2 == 0 else "F3F6F8")
            color = RED if str(value).upper() == "CONFIRM" else NAVY
            bold = idx == 0 or str(value).upper() == "CONFIRM"
            set_cell_text(cells[idx], value, color, bold, font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, title, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title.upper() if level == 1 else title)
    r.font.name = "Aptos Display"
    r.font.size = Pt(16 if level == 1 else 11)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD if level == 1 else NAVY)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), GOLD)
        pbdr.append(bottom)
        pPr.append(pbdr)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        text = text[len(bold_prefix):]
    r = p.add_run(text)
    r.font.name = "Aptos"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    return p


def add_callout(doc, title, text, color=PALE_GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.05)
    cell = table.cell(0, 0)
    shade(cell, color)
    margins(cell, 110, 130, 110, 130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    r2.font.name = "Aptos"
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.name = "Aptos"
        r.font.size = Pt(8.7)
        r.font.color.rgb = RGBColor.from_string(NAVY)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.58)
section.bottom_margin = Inches(0.58)
section.left_margin = Inches(0.62)
section.right_margin = Inches(0.62)
section.header_distance = Inches(0.25)
section.footer_distance = Inches(0.25)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9)
styles["Normal"].font.color.rgb = RGBColor.from_string(NAVY)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("PI PROP INTELLIGENCE  /  OWNER OPERATIONS")
hr.bold = True
hr.font.name = "Aptos"
hr.font.size = Pt(7.5)
hr.font.color.rgb = RGBColor.from_string(GOLD)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("CONFIDENTIAL  |  OPERATING COSTS & RENEWAL CALENDAR  |  UPDATED AUGUST 25, 2026")
fr.font.name = "Aptos"
fr.font.size = Pt(7)
fr.font.color.rgb = RGBColor.from_string(MID)

# Cover
cover = doc.add_table(rows=1, cols=1)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cover.autofit = False
cover.columns[0].width = Inches(7.05)
cell = cover.cell(0, 0)
shade(cell, NAVY)
margins(cell, 500, 300, 500, 300)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PI")
r.bold = True
r.font.name = "Aptos Display"
r.font.size = Pt(31)
r.font.color.rgb = RGBColor.from_string(GOLD)
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(8)
r2 = p2.add_run("PROP INTELLIGENCE")
r2.bold = True
r2.font.name = "Aptos Display"
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor.from_string(WHITE)
p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("OPERATING COSTS & RENEWAL CALENDAR")
r3.bold = True
r3.font.name = "Aptos"
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor.from_string(GOLD)
p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Owner reference  |  Fiscal year 2026-2027")
r4.font.name = "Aptos"
r4.font.size = Pt(9)
r4.font.color.rgb = RGBColor.from_string(LIGHT)

doc.add_paragraph()
add_callout(
    doc,
    "Important: complete the highlighted fields",
    "The application configuration identifies the services required to operate PI Prop Intelligence, but private invoice amounts and renewal dates are not stored in source control. Every CONFIRM field must be completed from the latest vendor invoice or billing dashboard. Never place passwords, API keys, card numbers, or service-role secrets in this document.",
)
add_heading(doc, "At-a-glance budget control", 2)
summary_rows = [
    ("Confirmed monthly recurring total", "$________________"),
    ("Confirmed annual recurring total", "$________________"),
    ("Estimated monthly usage charges", "$________________"),
    ("Annualized operating cost", "(Monthly x 12) + Annual + Usage"),
    ("Recommended operating contingency", "10-15% of annualized cost"),
    ("Next owner billing review", "________________"),
]
add_table(doc, ["CONTROL", "OWNER ENTRY"], summary_rows, [3.55, 3.50], 8.5)
add_body(doc, "Scope: production application, public marketing site, authentication, email, subscriptions, CI/CD, push notifications, sports data, domains, and supporting owner tools.", "Scope:")

doc.add_page_break()
add_heading(doc, "1. Critical production services")
critical = [
    ("Render", "Backend API, workers, database/cache connectivity", "Monthly / usage", "CONFIRM", "CONFIRM", "Critical", "App data and API stop"),
    ("Vercel", "Flutter web/PWA and marketing deployments", "Monthly or annual", "CONFIRM", "CONFIRM", "Critical", "Web properties stop updating"),
    ("Supabase", "Authentication, user data, database services", "Monthly / usage", "CONFIRM", "CONFIRM", "Critical", "Sign-in and user workflows fail"),
    ("Cloudflare", "DNS, CDN, domain routing and security", "Monthly / annual", "CONFIRM", "CONFIRM", "Critical", "Domains may not resolve"),
    ("Domain registrar", "propsintell.com and pipropsintell.com", "Annual", "CONFIRM", "CONFIRM", "Critical", "Website/email domains expire"),
    ("GitHub", "Source control and deployment integration", "Monthly / annual / free", "CONFIRM", "CONFIRM", "Critical", "Build and release workflow is impaired"),
    ("Codemagic", "iOS/Android CI builds and App Store upload", "Usage / monthly", "CONFIRM", "CONFIRM", "Critical", "No new mobile releases"),
    ("Apple Developer Program", "iOS signing, TestFlight, App Store distribution", "Annual", "CONFIRM", "CONFIRM", "Critical", "Certificates/distribution lapse"),
]
add_table(doc, ["SERVICE", "PURPOSE", "CADENCE", "AMOUNT", "NEXT DUE", "LEVEL", "IF INTERRUPTED"], critical, [1.0, 1.55, 0.78, 0.65, 0.72, 0.55, 1.8], 6.9)

add_heading(doc, "2. Customer access, messaging and subscriptions")
customer = [
    ("RevenueCat", "Subscription entitlements and receipt status", "Monthly / usage", "CONFIRM", "CONFIRM", "Critical"),
    ("App Store Connect", "iOS products, sales, reviews and payouts", "Commission / no fixed bill", "See Apple terms", "N/A", "Critical"),
    ("OneSignal", "Push notifications and notification service extension", "Monthly / usage", "CONFIRM", "CONFIRM", "Important"),
    ("Resend", "Transactional SMTP and authenticated email delivery", "Monthly / usage", "CONFIRM", "CONFIRM", "Critical"),
]
add_table(doc, ["SERVICE", "PURPOSE", "CADENCE", "AMOUNT", "NEXT DUE", "LEVEL"], customer, [1.15, 2.4, 1.15, 0.85, 0.85, 0.65], 7.4)
add_callout(doc, "Revenue products are not expenses", "Core Monthly ($29.99) and Core Annual ($299.99) are customer subscription products. Track gross receipts, Apple commission, refunds, taxes, and RevenueCat costs separately from the operating-expense total.", "DCEFEA")

doc.add_page_break()
add_heading(doc, "3. Sports data and content providers")
providers = [
    ("SportsGameOdds", "Primary sports odds/props and supplemental event coverage", "Active; two production keys", "Monthly / usage", "CONFIRM", "CONFIRM", "Critical"),
    ("The Odds API", "Secondary odds and market coverage", "Configured", "Monthly / usage", "CONFIRM", "CONFIRM", "Important"),
    ("API-Sports", "League/event data and supplemental baseball coverage", "Configured", "Monthly / usage", "CONFIRM", "CONFIRM", "Important"),
    ("SportsDataIO", "Supplemental sports data", "Configured", "Monthly / usage", "CONFIRM", "CONFIRM", "Important"),
    ("BALLDONTLIE", "Supplemental basketball data", "Configured", "Monthly / usage / free", "CONFIRM", "CONFIRM", "Supporting"),
    ("Sportradar", "WNBA and supplemental sports feeds", "Configured", "Contract / usage", "CONFIRM", "CONFIRM", "Supporting"),
    ("ESPN public data", "Schedules and player/headshot fallbacks", "Public source", "Free", "$0", "N/A", "Supporting"),
    ("MLB Stats API", "MLB schedules and player/headshot fallbacks", "Public source", "Free", "$0", "N/A", "Supporting"),
    ("Sportmonks", "Possible soccer supplement; verify whether still active", "VERIFY", "CONFIRM", "CONFIRM", "CONFIRM", "Optional"),
]
add_table(doc, ["PROVIDER", "ROLE", "STATUS", "CADENCE", "AMOUNT", "NEXT DUE", "LEVEL"], providers, [1.05, 2.05, 1.05, 0.9, 0.65, 0.72, 0.63], 6.9)
add_callout(doc, "Data-feed control", "For every paid feed, record the plan limit, request quota, overage rate, billing contact, and cancellation notice period in the vendor dashboard. Keep API keys only in approved secret stores such as Render, Vercel, Codemagic, or the designated password manager.")

add_heading(doc, "4. Optional owner productivity services")
optional = [
    ("ChatGPT / Codex", "Development, documentation, troubleshooting", "Not required at runtime", "CONFIRM", "CONFIRM"),
    ("Google Workspace / iCloud", "Owner communication and account recovery", "Verify active plan", "CONFIRM", "CONFIRM"),
    ("Discord", "Community or internal notifications if enabled", "Free / optional paid", "CONFIRM", "CONFIRM"),
    ("Squarespace", "Legacy site/domain service only if still retained", "VERIFY", "CONFIRM", "CONFIRM"),
]
add_table(doc, ["SERVICE", "PURPOSE", "STATUS", "AMOUNT", "NEXT DUE"], optional, [1.35, 2.4, 1.35, 0.95, 1.0], 7.6)

doc.add_page_break()
add_heading(doc, "5. Twelve-month renewal calendar")
calendar_rows = []
for month in ("September 2026", "October 2026", "November 2026", "December 2026", "January 2027", "February 2027", "March 2027", "April 2027", "May 2027", "June 2027", "July 2027", "August 2027"):
    calendar_rows.append((month, "________________", "________________", "$________", "[  ] Paid  [  ] Reviewed"))
add_table(doc, ["MONTH", "SERVICE / INVOICE", "DUE DATE", "AMOUNT", "STATUS"], calendar_rows, [1.3, 2.2, 1.15, 0.9, 1.5], 8.0)

add_heading(doc, "6. Monthly owner billing routine")
add_bullets(doc, [
    "On the first business day, review every vendor dashboard and reconcile charges against the business account.",
    "Confirm domain expiration dates, Apple membership status, payment cards, and billing-contact email access.",
    "Review usage limits for Render, Supabase, Vercel, Codemagic, RevenueCat, OneSignal, Resend, and each sports-data provider.",
    "Investigate unusual increases, failed payments, quota warnings, stale feeds, and services that are no longer used.",
    "Export invoices and receipts to the secure accounting folder using YYYY-MM - Vendor - Amount naming.",
    "Update this document after every price, plan, payment method, renewal date, or vendor change.",
])

add_heading(doc, "7. Immediate completion checklist")
checklist = [
    ("[  ]", "Enter the exact charge and next billing date for every CONFIRM field."),
    ("[  ]", "Record the legal billing account name and billing contact for each critical service."),
    ("[  ]", "Set 30-day and 7-day calendar reminders for annual renewals and domain expirations."),
    ("[  ]", "Confirm backup payment methods for Apple, domains, hosting, authentication, and email."),
    ("[  ]", "Verify that all invoices are delivered to an owner-controlled email account."),
    ("[  ]", "Cancel duplicate or legacy services only after confirming production does not depend on them."),
]
add_table(doc, ["DONE", "OWNER ACTION"], checklist, [0.65, 6.4], 8.2)

doc.add_page_break()
add_heading(doc, "8. Vendor record worksheet")
add_body(doc, "Duplicate this worksheet for any service requiring contract details beyond the master tables.")
worksheet = [
    ("Vendor / service", "____________________________________________"),
    ("Plan / tier", "____________________________________________"),
    ("Billing cadence", "____________________________________________"),
    ("Current amount", "$___________________________________________"),
    ("Next due / renewal date", "____________________________________________"),
    ("Auto-renewal", "[  ] Yes   [  ] No   [  ] Unknown"),
    ("Billing contact", "____________________________________________"),
    ("Cancellation notice", "____________________________________________"),
    ("Usage limit / overage", "____________________________________________"),
    ("Invoice storage location", "____________________________________________"),
    ("Last reviewed", "____________________________________________"),
    ("Owner notes", "____________________________________________\n____________________________________________"),
]
add_table(doc, ["FIELD", "OWNER ENTRY"], worksheet, [2.2, 4.85], 8.5)
add_callout(doc, "Security boundary", "Do not record passwords, API keys, private signing keys, Supabase service-role keys, Apple .p8 files, or full payment-card details here. Store secrets only in approved encrypted secret managers and platform environment settings.", "F7DDDD")

add_heading(doc, "9. Document control")
control = [
    ("Document owner", "PI Prop Intelligence owner"),
    ("Review frequency", "Monthly and after every vendor or pricing change"),
    ("Source of truth", "Latest vendor invoice and billing dashboard"),
    ("Prepared", "August 25, 2026"),
    ("Next scheduled review", "________________"),
    ("Version", "1.0"),
]
add_table(doc, ["CONTROL", "VALUE"], control, [2.2, 4.85], 8.5)
add_body(doc, "This tracker supports operational planning and is not accounting, tax, or legal advice. Reconcile all charges with your bookkeeper or accountant.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
