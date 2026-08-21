from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "updated_operations_guides"
NAVY = "071621"
NAVY_2 = "0D2535"
GOLD = "B59B52"
GOLD_LIGHT = "E5D39A"
SILVER = "D6DCE2"
MUTED = "667584"
WHITE = "FFFFFF"
BLACK = "17212B"
PALE = "F3F6F8"
GREEN = "198E72"
RED = "B9473C"
PAGE_WIDTH = 9360


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = qn(f"w:{edge}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)


def set_font(run, size=10.5, color=BLACK, bold=False, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def setup(title, short_title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(.72)
    section.bottom_margin = Inches(.68)
    section.left_margin = Inches(.82)
    section.right_margin = Inches(.82)
    section.header_distance = Inches(.28)
    section.footer_distance = Inches(.28)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for style_name, size, before, after in (
        ("Heading 1", 16, 17, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY_2 if style_name != "Heading 3" else GOLD)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run(f"PI PROP INTELLIGENCE  |  {short_title}"), 8.5, MUTED, True)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("PI PROP INTELLIGENCE LLC  |  AUGUST 2026  |  "), 8, MUTED)
    page_field(p)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("PI PROP INTELLIGENCE"), 11, GOLD, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(title), 28, NAVY, True, name="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    set_font(p.add_run(subtitle), 13, MUTED, False, True)
    add_callout(doc, "CURRENT EDITION", "Updated through August 21, 2026. Web/PWA release; no App Store or Play Store distribution is assumed.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    set_font(p.add_run("Prepared for PI PROP INTELLIGENCE LLC"), 10, NAVY_2, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Independent sports research and intelligence only. PI does not facilitate wagering or accept bets."), 8.5, MUTED)
    doc.add_page_break()
    return doc


def add_callout(doc, label, text, color=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [PAGE_WIDTH])
    shade(table.cell(0, 0), NAVY)
    p = table.cell(0, 0).paragraphs[0]
    set_font(p.add_run(label + "  "), 9, color, True)
    set_font(p.add_run(text), 9.5, WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None):
    if widths is None:
        widths = [PAGE_WIDTH // len(headers)] * len(headers)
        widths[-1] += PAGE_WIDTH - sum(widths)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, header in enumerate(headers):
        shade(hdr.cells[i], NAVY)
        p = hdr.cells[i].paragraphs[0]
        set_font(p.add_run(str(header)), 9, GOLD_LIGHT, True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2:
            for cell in cells:
                shade(cell, PALE)
        for i, value in enumerate(values):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            set_font(p.add_run(str(value)), 9.1, BLACK, i == 0)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), 10.5, BLACK, True)
        set_font(p.add_run(text[len(bold_lead):]), 10.5, BLACK)
    else:
        set_font(p.add_run(text), 10.5, BLACK)
    return p


def bullets(doc, items, numbered=False):
    if numbered:
        for index, item in enumerate(items, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.24)
            p.paragraph_format.first_line_indent = Inches(-0.24)
            set_font(p.add_run(f"{index}. "), 10.2, BLACK, bold=True)
            set_font(p.add_run(item), 10.2, BLACK)
        return
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(item), 10.2, BLACK)


def toc(doc, entries):
    h(doc, "Contents", 1)
    for number, title in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(f"{number}. "), 10, GOLD, True)
        set_font(p.add_run(title), 10, NAVY_2, True)
    doc.add_page_break()


def save(doc, filename):
    path = OUT / filename
    doc.core_properties.title = filename.replace("_", " ").replace(".docx", "")
    doc.core_properties.subject = "PI PROP INTELLIGENCE operations documentation"
    doc.core_properties.author = "PI PROP INTELLIGENCE LLC"
    doc.core_properties.keywords = "PROP INTELLIGENCE, operations, sports research, user guide"
    doc.save(path)
    print(path)


SPORTS = [
    ("MLB", "Baseball player props, scoreboard, game markets, Strikeout Pro Gold, and MLB umpire context."),
    ("NFL", "Football player props, scoreboard, game markets, injuries, line movement, and research tools."),
    ("NBA", "Basketball props, scoreboard, game markets, injuries, and referee tendencies."),
    ("WNBA", "Basketball props, scoreboard, game markets, injuries, and referee tendencies."),
    ("NHL", "Hockey player props, scoreboard, game markets, and research tools."),
    ("Soccer", "Supported competition props, scoreboard/game context, and market research; availability varies by league."),
    ("NCAAF", "College football props and scoreboard/game context when scheduled markets are available."),
    ("NCAAB", "College basketball props and scoreboard/game context when scheduled markets are available."),
    ("CFL", "Canadian football props and game context when provider inventory is available."),
]
PROP_SITES = [
    ("PrizePicks", "Pick'em prop inventory. PI displays provider lines and tracks research selections; confirm rules in PrizePicks."),
    ("Underdog", "Fantasy pick'em inventory and tracked slip presentation; provider rules and multipliers remain authoritative."),
    ("FanDuel", "Sportsbook prop and game-market inventory where licensed data is available; verify price/odds in FanDuel."),
    ("DraftKings Pick6", "Pick6 prop inventory, separate from DraftKings sportsbook lines."),
    ("Draft Picks", "Normalized provider label used when this feed returns inventory; do not confuse it with DraftKings."),
    ("Betr", "Betr Picks inventory where the connected feed provides active markets."),
]


def complete_guide():
    doc = setup("COMPLETE APPLICATION GUIDE", "APPLICATION GUIDE", "Every screen, control, workflow, status, provider, and operating rule")
    entries = [(i, title) for i, title in enumerate([
        "Product purpose and quick start", "Sign-in, account, and membership", "Workspace navigation", "Market Board", "Prop cards and research", "Sports and season status", "Scoreboard and Game Markets", "Research features", "Build and active slips", "Live tracking and history", "Prop Chat and notifications", "Supported sports and prop sites", "Owner tools", "Loading, empty, stale, and error states", "Troubleshooting and responsible use"
    ], 1)]
    toc(doc, entries)
    h(doc, "1. Product purpose and quick start")
    para(doc, "PI PROP INTELLIGENCE is a web-first sports research workspace. It organizes live prop markets, provider comparisons, model context, tracking, history, and community tools. It does not place wagers, hold funds, or replace the rules and settlement decisions of a sportsbook or fantasy operator.")
    bullets(doc, [
        "Open app.propsintell.com in a current browser and sign in.",
        "Select an in-season sport in the top banner. The app may choose an active MLB, NFL, NBA, or WNBA market at startup.",
        "Select a prop site or keep All Prop Sites, then review freshness before evaluating a card.",
        "Open View Research for a readable translucent research overlay; use the close control or click outside it to return.",
        "Select Over or Under to add a leg. Select the same side again or use Remove to undo it.",
        "Open the right-side Slip control to review, remove, save, or lock the active ticket.",
        "Use Slip Watcher for unresolved tickets and Past Slip History for settled tickets.",
    ], numbered=True)
    add_callout(doc, "VERIFY THE SOURCE", "Lines, odds, multipliers, market rules, and settlement terms can change. Always confirm the current value at the named prop site before making an external decision.")

    h(doc, "2. Sign-in, account, and membership")
    add_table(doc, ["Control", "Purpose", "How to use"], [
        ("Email / Password", "Authenticate an existing member", "Enter credentials and complete any requested CAPTCHA or email verification."),
        ("Continue with Google", "Use supported federated sign-in", "Choose the intended Google identity. PI creates a privacy-safe display name when needed."),
        ("Forgot Password", "Secure recovery", "Request the email, open the time-limited link, and set a new password."),
        ("Account rail", "Identity, plan, and billing", "Open Account to review membership, manage or restore access, and sign out."),
        ("Membership badge", "Shows Core, Pro, or Owner status", "The badge is informational. Server authorization, not color alone, controls access."),
        ("Sign Out", "Ends the current session", "Use the gold Sign Out control on shared devices."),
    ], [1650, 2600, 5110])
    para(doc, "Core provides standard research, board, tracking, history, and community access. Pro unlocks advanced model direction and protected intelligence tools. Owner access includes Pro capability plus private operations functions. Internal code may use the term Edge for the Pro entitlement; customer-facing documentation uses Pro.")

    h(doc, "3. Workspace navigation")
    add_table(doc, ["Area", "Items", "Behavior"], [
        ("Top sports banner", "Scoreboard, Sports, MLB, NFL, NBA, WNBA, NHL, Soccer, NCAAF, NCAAB, CFL, Owner Ops", "Sports stay across the top. Gold marks the active sport; inactive controls remain silver."),
        ("Research", "Today's Briefing, Market Board, Search Players, Line Movement, Injury Impact, Analytics, Prop Chat, Score Watch, The Lab, Referee/Officiating Tracker, EV Scanner", "Sections remain open until the user collapses them."),
        ("Build", "Prop Builder, Build Performance, Slip Watcher", "Build and organize research-backed tickets."),
        ("History", "Past Slip History, Track Record", "Review settled tickets and the published model record."),
        ("Specialty", "Strikeout Pro Gold", "MLB pitcher strikeout research for Pro members."),
        ("Right utility rail", "Account, Slip, live prop count, membership", "Labels remain visible; counters do not overlap icons."),
    ], [1300, 4650, 3410])

    h(doc, "4. Market Board")
    bullets(doc, [
        "Search is the widest control and matches player, team, or market text.",
        "Filters opens sport-aware category, side, tier, confidence, verdict, and other supported filters.",
        "Prop Chat opens the community workspace without changing the current market data.",
        "Provider chips show the first available sites; More opens additional sites. Gold marks only the selected provider.",
        "Status groups show Live/Stale, updated age, event count, active providers, and providers with no inventory.",
        "Category rails show All and only categories with current inventory. Counts distinguish total inventory from playable research candidates.",
        "Verdict groups such as Best Now, Slight Edge, Better Line, Wait/Monitor, and All Props organize recommendation readiness.",
        "Provider sections keep each site's props together. Balanced rows expand to fill the section without blank card slots.",
    ])
    add_table(doc, ["Board state", "Meaning", "User action"], [
        ("Live", "Current provider data is within freshness policy.", "Continue research and confirm the named source line."),
        ("Refresh Due / Stale Data", "The latest successful provider update is older than expected.", "Wait for recovery or refresh once; do not rely on a stale Pro recommendation."),
        ("No Inventory", "A connected provider returned no active props.", "This is not necessarily an outage; select another site or sport."),
        ("Recovery Running", "Automatic provider recovery is active.", "Avoid repeated manual refreshes while cooldown/retry logic runs."),
        ("Owner Alert", "Expected provider freshness crossed an owner threshold.", "Owner reviews provider, quota, workers, and last successful update."),
    ], [1850, 3650, 3860])

    h(doc, "5. Prop cards and research")
    add_table(doc, ["Card element", "What it means"], [
        ("Headshot / initials", "Stable keyed player identity. The previous valid image or branded fallback remains visible while a new image loads."),
        ("PI Pick / PI Signal / Model Pick", "The qualified system direction and selected line when authorized. A lean is not the same as a released play."),
        ("PI Trust", "Data-readiness and confidence context, not a guaranteed win rate."),
        ("Projection / verdict side", "Model estimate or research direction used in the explanation."),
        ("Under / Line / Over", "Select a side to add it; select the same side again or use Remove to undo."),
        ("EP", "Opens other active props for that player. Other card areas should not trigger this player-props overlay."),
        ("View Research", "Opens a blurred translucent modal with enlarged text, evidence chips, injury context, explainability, and research tools."),
        ("Provider and sport strip", "Identifies exactly which site and sport supplied the displayed market."),
    ], [2150, 7210])
    para(doc, "Research opens as a dedicated scrollable overlay rather than expanding between grid cards. Close Research dismisses the overlay. This protects card alignment and keeps detailed methodology readable on laptops and phones.")

    h(doc, "6. Sports and season status")
    add_table(doc, ["Situation", "Panel heading", "Behavior"], [
        ("Filters removed all results", "No Props Match These Filters", "Clear or adjust filters; sorting alone is not treated as a filter."),
        ("Upcoming games, markets not open", "Preseason - Markets Opening Soon", "Shows first scheduled game, countdown, and up to three matchups."),
        ("Game imminent, no props", "Markets Currently Unavailable", "The app keeps checking connected sites."),
        ("No future seasonal schedule", "Offseason", "Updates automatically when the scoreboard schedule becomes available."),
        ("Schedule lookup failed", "Schedule Temporarily Unavailable", "Does not mislabel a network failure as offseason."),
        ("Soccer or tennis gap", "No Upcoming Markets / Events Coming Soon", "Avoids inaccurate season language for multi-league or tour sports."),
    ], [2200, 2860, 4300])
    para(doc, "View Schedule opens the first available matchups. Notify Me enables session-level market notification feedback. Check Again reruns the prop query. Dates come from scoreboard data and are not hard-coded.")

    h(doc, "7. Scoreboard and Game Markets")
    bullets(doc, [
        "Scoreboard: choose date, status, and sport; review upcoming, live, and final games in local 12-hour time.",
        "Watch: add a game to Score Watch for score-change, final, overtime, or extended-play alerts.",
        "Game Markets: compare moneyline, spread, and total markets across available books.",
        "Refresh: requests the newest scoreboard state; authoritative final status is required before grading.",
        "The desktop left sidebar remains consistent when Scoreboard opens; user-collapsed groups stay collapsed until reopened.",
    ])

    h(doc, "8. Research features")
    add_table(doc, ["Feature", "Access", "How to use / what it does"], [
        ("Today's Briefing", "Public/Core", "Summarizes what clears the daily bar and explicitly states caveats or a no-play slate."),
        ("Search Players", "Public/Core", "Search a player, open focused research, and compare active markets and sites."),
        ("Line Movement", "Core", "Uses persisted opening observations and current lines; compare direction, magnitude, freshness, and market median."),
        ("Injury Impact", "Pro", "Reviews verified status, role, usage, opportunity, teammate context, and recheck guidance."),
        ("Analytics", "Core", "Filters historical results; use meaningful sample sizes rather than short streaks."),
        ("EV Scanner", "Pro", "Ranks qualified candidates by estimated EV, fair probability, or edge; confirm the live line."),
        ("Intelligence Lab", "Pro", "Runs correlation, scenario, matchup, fatigue, similarity, and simulation workflows."),
        ("Officiating Tracker", "Pro", "NBA/WNBA referee and MLB umpire profiles. NFL officials are not yet supported."),
        ("Prop Alerts", "Pro", "Displays triggered conditions and links back to the affected market for confirmation."),
        ("Strikeout Pro Gold", "Pro", "Ranks MLB pitcher strikeout opportunities with projection, edge, trust, lineup, weather, park, and umpire inputs."),
    ], [1850, 1050, 6460])

    h(doc, "9. Build and active slips")
    bullets(doc, [
        "Add a leg from a card by selecting Over or Under.",
        "The active Slip counter reflects selected legs and stays separate from the slip icon.",
        "Remove a leg by selecting its chosen side again or using the dedicated Remove control.",
        "Review provider, player, market, side, line, odds/multiplier context, and start time before locking.",
        "Same-provider rules and duplicate-market protection prevent incompatible or repeated legs.",
        "Prop Builder can apply sport, site, category, minimum trust/edge, and leg-count criteria, then allows removal or replacement before save.",
        "PI tracks research tickets only. Saving a slip in PI does not submit anything to a sportsbook or fantasy operator.",
    ], numbered=True)

    h(doc, "10. Live tracking and history")
    add_table(doc, ["Feature", "Purpose", "Key interpretation"], [
        ("Slip Watcher", "Unresolved locked tickets", "Shows live value, game state, and grading progress; settled tickets move to history."),
        ("Past Slip History", "Resolved tickets", "Filter All/Won/Lost and review ticket cards, leg results, and today's win/loss percentage."),
        ("Track Record", "Published model outcomes", "Rates remain contextualized by sample size; losses remain visible."),
        ("Build Performance", "Strategy review", "Compare ticket win rate with individual-leg hit rate by sport, site, player, and market."),
        ("Closing-line value", "Entry versus close", "Positive CLV is useful process evidence but does not guarantee the result."),
    ], [1800, 3000, 4560])

    h(doc, "11. Prop Chat and notifications")
    bullets(doc, [
        "Use a public username; never post passwords, payment details, private contact information, or account tokens.",
        "Community chat supports text, approved HTTPS links, attachments, shared props, and shared slips.",
        "Direct contacts, mentions, and unread counts use gold badges. Owner messages and identity indicators are visibly distinct.",
        "The floating chat bubble can be dismissed and restored. It should not cover key utility controls.",
        "Push and sound preferences are device-specific. Browser permission must be granted before delivery can begin.",
    ])

    h(doc, "12. Supported sports and prop sites")
    add_table(doc, ["Sport", "Current application use"], SPORTS, [1200, 8160])
    add_table(doc, ["Prop site", "How PI uses it"], PROP_SITES, [1900, 7460])
    add_callout(doc, "SITE DISTINCTION", "Prop sites are customer-visible market sources. The Odds API, SportsGameOdds, API-Sports, ESPN, Sportmonks, Supabase, Render, Vercel, RevenueCat, OneSignal, and Cloudflare are infrastructure or data services, not selectable prop sites.")

    h(doc, "13. Owner tools")
    bullets(doc, [
        "Owner Ops is available only to the verified owner role. Paid tier alone does not grant access.",
        "Run All Checks refreshes API, database, Redis, queue/worker, provider, quota, freshness, scoreboard, billing, active-user, grading, and deployment signals.",
        "Provider status distinguishes Active, No Current Inventory, Stale, and Outage. Last successful update is shown when available.",
        "Owner top picks surface up to four qualified picks per supported sport for owner use; normal release gates and data caveats still apply.",
        "Issues Requiring Review contains overdue pending, questionable grading, provider, pipeline, and billing follow-ups.",
        "Owner tier preview changes visible member experience without changing the owner's protected authorization.",
    ])

    h(doc, "14. Loading, empty, stale, and error states")
    add_table(doc, ["State", "What the interface should show", "What not to assume"], [
        ("Loading", "Stable skeleton cards and branded headshot fallbacks", "A blank circle or moving layout is not an acceptable loading state."),
        ("Empty", "Filter-specific or schedule-aware explanation", "Zero cards does not automatically mean provider failure."),
        ("Stale", "Age, affected providers, and recovery guidance", "Old inventory is not current merely because it remains visible."),
        ("Error", "Plain-language message and safe retry", "Raw URLs, credentials, stack traces, or private provider details must never be shown."),
        ("Offline", "Connection-specific recovery message", "Do not repeatedly refresh if the device has no network."),
    ], [1450, 4350, 3560])

    h(doc, "15. Troubleshooting and responsible use")
    add_table(doc, ["Problem", "Member action", "Owner action"], [
        ("No props", "Check selected sport/site and season panel; clear real filters.", "Check provider inventory, freshness, quota, and event schedule."),
        ("Photos missing", "Wait for branded initials/fallback; refresh once.", "Check headshot mappings, cache sync, and provider URLs."),
        ("Recommendation withheld", "Read the explanation and treat the line as information only.", "Confirm identity, freshness, model sample, projection, and release gate."),
        ("Cannot sign in", "Use recovery and confirm the intended identity.", "Check Supabase Auth, Turnstile, redirects, and rate limits."),
        ("Paid access missing", "Refresh Account and restore/manage subscription.", "Check RevenueCat entitlement, webhook, and Supabase tier mapping."),
        ("Slip pending", "Wait for official final status and stats.", "Inspect grading review, event linkage, and authoritative results."),
    ], [1550, 3900, 3910])
    add_callout(doc, "RESPONSIBLE USE", "PI outputs are research estimates, not guarantees. Verify live lines, rules, eligibility, and settlement terms with the external operator. Never risk funds needed for essential expenses.")
    return doc


def infrastructure_guide():
    doc = setup("INFRASTRUCTURE GUIDE", "INFRASTRUCTURE", "Domains, platforms, data sources, security boundaries, deployments, and recovery")
    toc(doc, [(i, title) for i, title in enumerate([
        "Architecture", "Production sites and dashboards", "Customer-facing prop sites", "Data and content providers", "Application services", "Data flow and freshness", "Security boundary", "Deployment and verification", "Monitoring and incident response", "Backup and recovery", "Known boundaries"
    ], 1)])
    h(doc, "1. Architecture")
    para(doc, "The production platform is a Flutter web/PWA client served through Vercel and a FastAPI/Python backend served through Render. Supabase provides authentication and PostgreSQL data services. Redis and RQ provide shared cache, queueing, coordination, rate limits, and short-lived operational state. External providers supply markets, scoreboards, statistics, images, billing, CAPTCHA, and notifications.")
    add_table(doc, ["Layer", "Service", "Responsibility"], [
        ("Public web", "propsintell.com", "Marketing, product, pricing, legal, installation, contact, and authentication entry points."),
        ("Application", "app.propsintell.com / Vercel", "Flutter web/PWA, responsive shell, browser-delivered public configuration, and release caching."),
        ("API", "api.propsintell.com / Render", "Authentication enforcement, prop catalog, models, grading, provider orchestration, webhooks, and owner telemetry."),
        ("Identity/data", "Supabase", "Authentication, profiles, PostgreSQL, row-level security, storage, and managed backups."),
        ("Cache/jobs", "Redis + RQ on Render", "Distributed catalog, queues, workers, cooldowns, rate limiting, image maps, and recovery state."),
        ("Source/release", "GitHub", "Repository, commit history, CI, release source of truth, and deployment triggers."),
    ], [1300, 2500, 5560])

    h(doc, "2. Production sites and dashboards")
    add_table(doc, ["Site", "Use", "Owner verification"], [
        ("Vercel", "Frontend deployment and public environment variables", "Latest main deployment is Ready; custom domain points to it; update prompt serves the intended bundle."),
        ("Render", "API, worker, Redis, cron, secrets, and logs", "Services Live; deploy commit matches GitHub; workers/cron healthy; no restart loop."),
        ("Supabase", "Auth, PostgreSQL, RLS, backups", "Owner MFA; redirects; Security Advisor; backups/PITR; no exposed service role."),
        ("GitHub", "Code, commits, Actions", "main traceable; checks understood; no secrets; deployment source matches intended commit."),
        ("RevenueCat", "Products, offerings, entitlements, webhook", "Core/Pro mapping valid; webhook authenticated; sandbox and production separated."),
        ("OneSignal", "Push delivery and device tags", "App ID and backend credentials configured; permission and test notification verified."),
        ("Cloudflare Turnstile", "Authentication bot protection", "Site key covers app domain; secret remains server-side/Supabase; hostname is correct."),
    ], [1750, 3420, 4190])

    h(doc, "3. Customer-facing prop sites")
    add_table(doc, ["Prop site", "Normalized label and behavior"], PROP_SITES, [1900, 7460])
    para(doc, "Provider separation is a presentation and trust requirement: props from different sites are grouped under separate provider headers and are not mixed into one indistinguishable card list. DraftKings Pick6 and sportsbook DraftKings are distinct products.")

    h(doc, "4. Data and content providers")
    add_table(doc, ["Provider", "Role", "Protection / fallback"], [
        ("The Odds API", "Sportsbook events, player props, game markets, quota telemetry", "Quota reserve, nearest-event priority, adaptive cooldown, cache, stale disclosure, and provider health."),
        ("SportsGameOdds", "Supplemental prop events and provider inventory", "Retry-After, exponential cooldown, health counters, and no request cascade during rate limits."),
        ("API-Sports", "Selected scoreboard, player statistics, and grading support", "Sport-aware fallback, cache, and explicit provider error reporting."),
        ("ESPN", "Scoreboard/stat context and player images", "Redis/bundled headshot maps and source-specific fallback."),
        ("Sportmonks", "Soccer history, player context, and images", "Shared cache plus local fallback; coverage depends on subscribed competition."),
        ("Official league APIs", "MLB and supported official context", "Cached scheduled ingestion; request-time paths stay bounded."),
    ], [1800, 3500, 4060])

    h(doc, "5. Application services")
    bullets(doc, [
        "Flutter/Dart: responsive board, scoreboards, research overlays, builders, history, chat, account, billing, and owner operations.",
        "FastAPI/Python: API routing, token validation, authorization, normalization, modeling, line history, grading, provider health, webhooks, and telemetry.",
        "PostgreSQL: user, profile, ticket, prediction, historical, officiating, line-snapshot, engagement, billing, and security-event records.",
        "Redis/RQ: compressed prop catalog, shared cache, queue coordination, retries, rate limits, provider cooldowns, and headshot maps.",
        "RevenueCat: Core and Pro entitlement state. Backend-verified webhook results control durable access.",
        "OneSignal: web push delivery. The browser permission and member/device association are separate prerequisites.",
    ])

    h(doc, "6. Data flow and freshness")
    bullets(doc, [
        "Scheduled/live sync jobs fetch provider inventory with bounded concurrency and quota controls.",
        "Provider payloads are normalized into one prop contract while preserving source site, event, player, market, and update time.",
        "The catalog assigns stable group identities and player-image keys before publishing compressed shared state.",
        "Every pregame line change is recorded in sportsbook_line_snapshots. The live catalog reapplies the earliest observation to expose genuine opening-to-current movement.",
        "The API filters by sport, site, category, verdict, freshness, and access tier; the frontend must not reconstruct protected Pro direction.",
        "Stale/no-inventory/outage states remain distinct. Owners see last successful provider updates and alerts when an expected source becomes stale.",
    ], numbered=True)
    add_table(doc, ["Provider state", "Definition", "Operational response"], [
        ("Active", "Fresh inventory returned", "Monitor normally."),
        ("No Current Inventory", "Provider healthy but no props matched the current schedule", "Do not classify as outage."),
        ("Stale", "Last successful inventory exceeds freshness threshold", "Inspect sync, quota, cooldown, cache, and worker."),
        ("Outage", "Expected provider failed or cannot be reached", "Use fallback/cached disclosure and open incident follow-up."),
    ], [1800, 3860, 3700])

    h(doc, "7. Security boundary")
    add_callout(doc, "NEVER IN THE BROWSER", "Provider secrets, Supabase service-role key, database connection strings, Redis URL, RevenueCat webhook authorization, private signing secrets, deployment tokens, or user access tokens in logs.", RED)
    bullets(doc, [
        "Allowed public configuration includes Supabase publishable/anon key, Turnstile site key, RevenueCat public SDK key, OneSignal public app ID, and API base URL.",
        "The backend validates Supabase tokens and independently enforces owner, admin, Core, and Pro access.",
        "Database row-level security remains enabled even when server routes also authorize access.",
        "CORS permits approved production origins. Rate limiting and safe errors protect public routes.",
        "Logs omit authorization headers, complete billing objects, secrets, and unnecessary personal data.",
        "Owner Ops is owner-only. Admin/tester access is not interchangeable with owner authorization.",
    ])

    h(doc, "8. Deployment and verification")
    bullets(doc, [
        "Review intended files and test results; exclude unrelated or generated artifacts.",
        "Commit to main with a descriptive message and push GitHub.",
        "Confirm Vercel and Render deploy the same intended commit.",
        "Verify public site and app return successfully; update the PWA when the release prompt appears.",
        "Verify /ready dependencies, prop catalog, auth protection, provider state, scoreboard, headshots, filters, active slip, research overlay, and owner operations.",
        "Record the commit, deployment timestamps, known warnings, and rollback candidate.",
    ], numbered=True)

    h(doc, "9. Monitoring and incident response")
    add_table(doc, ["Symptom", "First check", "Next action"], [
        ("App fails to load", "Vercel deployment/domain", "Inspect build, custom domain, public variables, and PWA cache."),
        ("App loads; props fail", "Owner Ops and Render API", "Check release, catalog source, providers, quota, Redis, queue, and freshness."),
        ("Stale provider alert", "Provider last success and cooldown", "Do not force-refresh through rate limit; verify recovery/fallback."),
        ("Photos fail", "Headshot cache/source", "Confirm stable keys, mappings, scheduled image sync, and fallback behavior."),
        ("Login/CAPTCHA fails", "Supabase Auth and Turnstile", "Check site/secret pairing, hostnames, redirects, and auth limits."),
        ("Paid access missing", "RevenueCat then Supabase", "Verify entitlement, webhook authentication/idempotency, and profile tier."),
        ("Tickets unsettled", "Owner grading review", "Confirm final game state, event identity, result source, and reconciliation."),
    ], [1800, 3000, 4560])

    h(doc, "10. Backup and recovery")
    bullets(doc, [
        "Confirm Supabase backup retention and point-in-time recovery policy monthly.",
        "Preserve durable ticket, prediction, line-history, and migration-checksum tables through code releases.",
        "Redis is operational state, not the only durable source for critical user history.",
        "Recover code by redeploying a known healthy commit; never delete production data to fix a frontend release.",
        "Test restore procedures and document ownership before an incident, not during one.",
    ])

    h(doc, "11. Known boundaries")
    bullets(doc, [
        "NFL officiating profiles are not currently implemented; Officiating Tracker supports NBA, WNBA, and MLB.",
        "Provider and sport inventory changes with schedule, subscription coverage, geography, and upstream market release timing.",
        "The web/PWA is the supported distribution path for this edition; App Store and Play Store release procedures are out of scope.",
        "No browser application can hide downloaded client code. Proprietary logic, secrets, and protected payloads remain server-side.",
    ])
    return doc


def owner_schedule():
    doc = setup("OWNER OPERATIONS SCHEDULE", "OWNER RUNBOOK", "Daily, weekly, monthly, seasonal, release, and incident procedures")
    toc(doc, [(i, title) for i, title in enumerate([
        "Daily opening", "Daily closing", "Weekly review", "Monthly review", "Season and provider readiness", "Release day", "Incident playbooks", "Feature smoke checklist", "Owner records and escalation"
    ], 1)])
    h(doc, "1. Daily opening check (10-15 minutes)")
    bullets(doc, [
        "Sign in with the verified owner account; confirm Pro/Owner mode and open Owner Ops.",
        "Run All Checks and wait for current timestamps.",
        "Confirm API, PostgreSQL, Redis, queue, and expected workers are healthy.",
        "Review provider state: Active, No Current Inventory, Stale, and Outage; confirm last successful update per expected provider.",
        "Check quota reserve, cooldown/rate-limit state, failed jobs, and scheduled sync completion.",
        "Confirm board freshness, event counts, playable counts, scoreboard latency, and image/headshot health.",
        "Review failed payments, unsettled slips, questionable grades, pipeline alerts, and owner notifications.",
        "Confirm production frontend/backend commit matches the intended release.",
        "Open Market Board, one prop research overlay, Scoreboard, Active Slip, and an empty-season sport on desktop and phone width.",
    ], numbered=True)
    add_table(doc, ["Signal", "Healthy target", "Escalate when"], [
        ("API / database / Redis", "Healthy and current", "Any dependency is unavailable or repeated recovery mode persists."),
        ("Providers", "Expected sources Active or honestly No Inventory", "Expected source is Stale/Outage or last success crosses policy."),
        ("Props", "Counts fit current event schedule", "Sudden unexplained collapse, wrong sport, mixed providers, or stale recommendations."),
        ("Images", "Headshots or branded fallback remain stable", "Black circles, flicker, repeated URL failures, or mapping collapse."),
        ("Grading", "No unexplained overdue pending legs", "Final games remain unresolved or manual correction lacks evidence."),
    ], [1900, 3650, 3810])

    h(doc, "2. Daily closing check (5-10 minutes)")
    bullets(doc, [
        "Review provider alerts and document any outage or low-quota period.",
        "Review today's locked/settled tickets and current win/loss percentage; investigate inconsistent grading.",
        "Confirm line snapshots recorded for upcoming markets and no stale recommendation remained promoted.",
        "Check failed payments, auth/security alerts, unresolved chat reports, and owner follow-ups.",
        "Record incidents, manual actions, and the next owner/date before signing out.",
    ], numbered=True)

    h(doc, "3. Weekly review (30-45 minutes)")
    add_table(doc, ["Area", "Review", "Expected result"], [
        ("Providers", "Error rate, quota, cooldowns, coverage, inventory gaps, last success", "No unexplained stale source; no outage mislabeled as no inventory."),
        ("Model / track record", "Sample, calibration, release gates, CLV, sport/site/market breakdown", "Rates shown with adequate sample; losses and caveats remain visible."),
        ("Workers / cron", "Queue depth, retries, failed jobs, historical/headshot/schedule jobs", "Expected jobs complete; no old failure backlog."),
        ("Billing", "Webhook authentication, idempotency, failed payments, entitlement mismatch", "Access matches verified RevenueCat state."),
        ("Auth / security", "Supabase advisor, Turnstile, rate limits, owner MFA, logs", "No exposed secrets or unresolved security finding."),
        ("Performance", "First load, board, scoreboard, overlays, mobile interactions, API latency", "No significant regression from the prior week."),
    ], [1600, 4340, 3420])

    h(doc, "4. Monthly review (60-90 minutes)")
    bullets(doc, [
        "Confirm Supabase backup and retention/PITR policy; document restore readiness.",
        "Review Vercel, Render, Redis, database, provider, RevenueCat, and OneSignal usage/cost trends.",
        "Audit owner/admin/tester roles, paid entitlements, inactive accounts, and public signup policy.",
        "Review privacy, terms, contact, pricing, and Complete Application Guide for changes.",
        "Reconcile supported sports and six prop-site labels against current configuration.",
        "Review model drift, calibration, release thresholds, stale-data suppression, and owner top-pick quality.",
        "Run desktop, narrow phone, tablet, laptop, and ultrawide visual checks; include browser zoom and keyboard navigation.",
    ])

    h(doc, "5. Season and provider readiness")
    add_table(doc, ["Timing", "Actions", "Verify"], [
        ("Before season", "Confirm schedule source, sport navigation, provider mappings, categories, player identity, headshots, and model inputs.", "Season panel shows real next event; no hard-coded date; markets transition automatically."),
        ("Market opening", "Confirm provider inventory, site separation, counts, line history, freshness, and playable classifications.", "Cards match source; no mixed provider rows or blank grid slots."),
        ("In season", "Monitor daily schedule, quota, late lineup/injury changes, stale alerts, and grading.", "Active data is current; owner alerts are actionable."),
        ("Postseason", "Confirm event filters and schedule boundaries; preserve playoff context.", "No premature offseason label."),
        ("Offseason", "Confirm schedule-aware empty state and historical research availability.", "No generic broken-feed message when the season is legitimately inactive."),
    ], [1500, 4930, 2930])

    h(doc, "6. Release day procedure")
    bullets(doc, [
        "Review worktree and intended files; leave unrelated artifacts uncommitted.",
        "Run targeted tests for changed features and the broad backend suite; understand any failing Flutter tests before release.",
        "Commit with a descriptive message and push main.",
        "Confirm GitHub, Vercel, and Render report the intended commit.",
        "Update the production PWA when prompted, then hard-refresh only if the old bundle remains.",
        "Smoke auth, board, sport/provider filters, headshots, research overlay, add/remove selection, active slip, scoreboard, history, chat, membership, and Owner Ops.",
        "Check console/frontend errors, API latency, image failures, provider freshness, and queue health for at least one full refresh cycle.",
        "Record release commit, result, known warnings, and rollback commit.",
    ], numbered=True)

    h(doc, "7. Incident playbooks")
    add_table(doc, ["Signal", "Immediate action", "Do not"], [
        ("API unhealthy", "Check Render deploy/restarts, /ready, database, Redis, and release commit.", "Do not redeploy repeatedly without isolating the dependency."),
        ("Provider stale/outage", "Check last success, quota, cooldown, worker, cache, and upstream response.", "Do not force-refresh through a rate limit or call no inventory an outage."),
        ("Prop count collapse", "Compare schedule, selected sport/site, provider totals, catalog source, and filters.", "Do not publish placeholder or fabricated props."),
        ("Headshot failures", "Check mapping key, image URL, cache sync, and branded fallback.", "Do not render an empty black circle."),
        ("Scoreboard slow", "Check provider latency, cache path, and API timing.", "Do not disable caching as the first response."),
        ("Paid access mismatch", "Check RevenueCat event, entitlement, webhook, and Supabase profile.", "Do not grant access from a browser success message alone."),
        ("Questionable grade", "Confirm final event, official stat, market rule, and stored selection.", "Do not manually grade from incomplete live data."),
    ], [1700, 4250, 3410])

    h(doc, "8. Feature smoke checklist")
    add_table(doc, ["Feature", "Minimum acceptance check"], [
        ("Market Board", "Props load; providers separated; balanced rows; filter/search clear correctly; status age is honest."),
        ("Season status", "Sport-only empty state shows schedule; real filters show filtered-empty; date/countdown derives from live schedule."),
        ("Prop card", "Headshot stable; EP alone opens player props; Over/Under add and undo; View Research opens readable overlay."),
        ("Line Movement", "At least known moved fixtures hydrate opening/current values; no false same-line movement."),
        ("Officiating", "NBA/WNBA referees and MLB umpires load or show a truthful no-data state; no NFL tab."),
        ("Slip", "Counter, icon, remove, save/lock, live tracking, and settlement transitions work."),
        ("History", "All/Won/Lost, today win/loss percentage, result cards, and text fit without clipping."),
        ("Responsive shell", "Left sections retain state; right rail labels/counters do not overlap; mobile filters/dialogs do not clip."),
    ], [1900, 7460])

    h(doc, "9. Owner records and escalation")
    add_table(doc, ["Record", "Value"], [
        ("Review date / owner", "____________________________________________"),
        ("Production commit", "____________________________________________"),
        ("Frontend/backend deploy status", "____________________________________________"),
        ("Provider freshness / quota", "____________________________________________"),
        ("Props / playable / events", "____________________________________________"),
        ("Model sample / calibration", "____________________________________________"),
        ("Unsettled / questionable items", "____________________________________________"),
        ("Backup / security verified", "____________________________________________"),
        ("Actions and follow-up owner/date", "____________________________________________"),
    ], [3300, 6060])
    return doc


def features_guide():
    doc = setup("CURRENT APPLICATION FEATURES", "FEATURE INVENTORY", "Authoritative web/PWA feature, access, provider, and limitation inventory")
    toc(doc, [(i, title) for i, title in enumerate([
        "Product access", "Supported sports and prop sites", "Navigation and workspace", "Market discovery", "Research and intelligence", "Scoreboards and game markets", "Build, slips, and history", "Chat and notifications", "Accounts and membership", "Owner operations", "Reliability and security", "Recent changes", "Current limitations"
    ], 1)])
    h(doc, "1. Product access")
    add_table(doc, ["Surface", "Address / delivery", "Capability"], [
        ("Public site", "propsintell.com", "Product, feature, pricing, terms, privacy, contact, install, and sign-in entry."),
        ("Application", "app.propsintell.com", "Responsive Flutter web/PWA for desktop, tablet, and mobile browsers."),
        ("Updates", "Automatic web release", "No App Store or Play Store download required; update prompt loads the newest bundle."),
        ("API", "api.propsintell.com", "Authenticated data, intelligence, tracking, billing, and owner operations."),
    ], [1500, 2600, 5260])

    h(doc, "2. Supported sports and prop sites")
    add_table(doc, ["Sport", "Feature coverage"], SPORTS, [1200, 8160])
    add_table(doc, ["Prop site", "Current behavior"], PROP_SITES, [1900, 7460])

    h(doc, "3. Navigation and workspace")
    add_table(doc, ["Navigation group", "Current destinations"], [
        ("Top banner", "Scoreboard, Sports, MLB, NFL, NBA, WNBA, NHL, Soccer, NCAAF, NCAAB, CFL, Owner Ops."),
        ("Research", "Today's Briefing, Market Board, Search Players, Line Movement, Injury Impact, Analytics, Prop Chat, Score Watch, The Lab, Referee/Officiating Tracker, EV Scanner."),
        ("Build", "Prop Builder, Build Performance, Slip Watcher."),
        ("History", "Past Slip History, Track Record."),
        ("Specialty / Owner", "Strikeout Pro Gold; Owner Operations Center for the verified owner."),
        ("Utility rail", "Account, active Slip, current live prop count, and membership badge."),
    ], [2200, 7160])

    h(doc, "4. Market discovery")
    bullets(doc, [
        "Live multi-sport prop catalog with provider, event, market, line, update, and freshness identity.",
        "All Prop Sites plus six normalized site filters; More menu for overflow.",
        "Provider-separated sections and balanced rows with no blank placeholder slots.",
        "Search, sport, provider, category, side, tier, confidence, verdict, and sort controls.",
        "Sport-aware category counts, playable counts, Best Now, Slight Edge, Better Line, Wait/Monitor, and All Props groups.",
        "Schedule-aware empty state with first event, countdown, upcoming games, market-opening guidance, session notifications, and historical research direction.",
        "Stable keyed headshots with previous-frame retention and branded initials/silhouette fallback.",
        "Over/Under selection, explicit undo/removal, favorite, and EP-only player-market overlay.",
        "Readable translucent View Research overlay with enlarged text and scrolling.",
    ])

    h(doc, "5. Research and intelligence")
    add_table(doc, ["Feature", "Minimum access", "Status / purpose"], [
        ("Today's Briefing", "Public/Core", "Daily qualified opportunities and explicit caveats/no-play outcome."),
        ("Search Players", "Public/Core", "Focused active market and provider comparison."),
        ("Analytics", "Core", "Result and performance filtering with sample context."),
        ("Line Movement", "Core", "Persisted opening/current line history, freshness, direction, and provider comparison."),
        ("Injury Impact", "Pro", "Verified availability, role, usage, opportunity, and teammate-context changes."),
        ("EV Scanner", "Pro", "Qualified estimated EV, fair probability, edge, and offered-price comparison."),
        ("Intelligence Lab", "Pro", "Correlation, scenarios, simulations, fatigue, matchup, and historical analogs."),
        ("Officiating Tracker", "Pro", "NBA/WNBA referee and MLB umpire tendencies; NFL not implemented."),
        ("Prop Alerts", "Pro", "Tracked condition and line-change alerts."),
        ("Strikeout Pro Gold", "Pro", "MLB strikeout model workflow with methodology/explainability."),
    ], [2000, 1300, 6060])

    h(doc, "6. Scoreboards and game markets")
    bullets(doc, [
        "Live, upcoming, and final scoreboards with date and sport filters.",
        "Local 12-hour time, event status, score, league, and supported game detail.",
        "Score Watch for selected games and score/final/extended-play alerts.",
        "Game Markets for moneyline, spread, total, and available book comparison.",
        "Scoreboard schedule powers empty-sport season/market status panels.",
    ])

    h(doc, "7. Build, slips, and history")
    bullets(doc, [
        "Manual Prop Builder with sport/site/category and minimum research thresholds.",
        "Same-provider and duplicate-market protections.",
        "Active Slip review with provider-specific presentation, add/remove, save/lock, and status.",
        "Slip Watcher live values, game status, grading, settlement, and automatic movement to history.",
        "Past Slip History All/Won/Lost filtering, readable ticket cards, today win/loss percentage, and historical result context.",
        "Track Record keeps model winners and losses visible with sample-size context.",
        "Build Performance analyzes ticket and leg outcomes by sport, site, player, and market.",
    ])

    h(doc, "8. Chat and notifications")
    bullets(doc, [
        "Community Prop Chat and privacy-safe usernames.",
        "Direct contacts, mentions, unread badges, HTTPS links, attachments, shared props, and shared slips.",
        "Owner identity/message treatment in gold and separate role authorization.",
        "Floating chat bubble with dismiss/restore behavior.",
        "OneSignal web push and persisted sound preference where browser permission exists.",
    ])

    h(doc, "9. Accounts and membership")
    add_table(doc, ["Capability", "Core", "Pro / Owner"], [
        ("Live board, search, provider comparison", "Included", "Included"),
        ("Scoreboard and Game Markets", "Included", "Included"),
        ("Manual builder, active slips, tracking", "Included", "Included / enhanced"),
        ("History and Track Record", "Standard/recent", "Full/advanced"),
        ("System Over/Under model direction", "Not shown", "Qualified signals only"),
        ("Projection, confidence, edge, advanced line context", "Factual/basic", "Included when qualified"),
        ("EV Scanner, Lab, Injury, Alerts, Officiating", "Not included", "Included"),
        ("Strikeout Pro Gold", "Not included", "Included"),
        ("Owner Operations", "Not included", "Verified owner only"),
    ], [3500, 2100, 3760])
    para(doc, "Customer-facing tiers are Core and Pro. The verified owner receives protected owner functions without weakening server-side role checks. Pricing displayed in the live application and billing platform is authoritative.")

    h(doc, "10. Owner operations")
    bullets(doc, [
        "Run All Checks production control panel.",
        "API, database, Redis, workers, queues, provider quality/quota/freshness, scoreboard latency, billing, grading, users, and deployment signals.",
        "Provider state distinguishes active inventory, healthy/no inventory, stale, and outage; last successful update and owner alerts are supported.",
        "Issues Requiring Review for questionable grades, unsettled tickets, pipelines, providers, and billing.",
        "Owner-only top four qualified picks per sport, subject to normal data and release gates.",
        "Owner tier preview and account-role management remain separate from customer billing tier.",
    ])

    h(doc, "11. Reliability and security")
    bullets(doc, [
        "Compressed distributed prop catalog plus durable catalog recovery.",
        "Background provider refresh, retries, cooldowns, quota reserve, and stale disclosure.",
        "Persistent line snapshots for movement and closing-line analysis.",
        "Stable skeletons, image fallback, no black headshot state, and reduced layout movement.",
        "Supabase authentication, PostgreSQL RLS, backend access enforcement, CORS, rate limits, and safe logs/errors.",
        "GitHub-driven release traceability with Vercel frontend and Render backend deployment.",
    ])

    h(doc, "12. Recent changes included in this edition")
    add_table(doc, ["Area", "Updated behavior"], [
        ("Navigation", "Sports across top; Research/Build/History on left; scoreboard retained; owner utility rail labels."),
        ("Providers", "Separated provider sections, More overflow, normalized sites, active/no-inventory/stale/outage status."),
        ("Cards", "Stable images, explicit selection undo, EP-only player overlay, research modal, no external provider strip."),
        ("Empty sports", "Schedule-derived market/preseason/offseason panel with real event countdown."),
        ("Line Movement", "Recorded snapshots hydrate opening/current values."),
        ("Officiating", "MLB umpires added to NBA/WNBA referee support."),
        ("History", "Readable result text and current win/loss percentage."),
        ("Responsive", "Balanced card rows, collapsible side groups, mobile density reductions, and clipping protections."),
    ], [1900, 7460])

    h(doc, "13. Current limitations")
    bullets(doc, [
        "Live inventory depends on provider schedule, subscription coverage, geography, market release, quota, and upstream health.",
        "NFL officiating ingestion is not yet available and is not shown as an empty tab.",
        "Headshot quality depends on mapping/source availability, but the interface should always retain a branded fallback.",
        "Web/PWA is the supported release. Native-store distribution is not part of the current launch plan.",
        "PI does not execute transactions, guarantee outcomes, or override an external operator's market and settlement rules.",
    ])
    return doc


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    save(complete_guide(), "PI_PROP_INTELLIGENCE_Complete_Application_Guide_August_2026.docx")
    save(infrastructure_guide(), "PI_PROP_INTELLIGENCE_Infrastructure_Guide_August_2026.docx")
    save(owner_schedule(), "PI_PROP_INTELLIGENCE_Owner_Operations_Schedule_August_2026.docx")
    save(features_guide(), "PROP_INTELLIGENCE_Current_Features_August_2026.docx")
